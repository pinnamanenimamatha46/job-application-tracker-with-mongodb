import os
from pathlib import Path

from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader

from app.schemas.resume_analysis import ResumeAnalysis


load_dotenv()

SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".docx", ".txt"}


def get_openai_client() -> OpenAI:
    """Create an OpenAI client using the environment API key."""
    api_key = os.getenv("OPENAI_API_KEY")

    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY is missing. Add it to the .env file."
        )

    return OpenAI(api_key=api_key)


def analyze_resume_match(
    resume_text: str,
    job_description: str,
) -> ResumeAnalysis:
    """Compare resume text with a job description using AI."""
    if not resume_text.strip():
        raise ValueError("Resume text cannot be empty.")

    if not job_description.strip():
        raise ValueError("Job description cannot be empty.")

    client = get_openai_client()

    response = client.responses.parse(
        model="gpt-4.1-mini",
        input=[
            {
                "role": "system",
                "content": (
                    "You are an expert technical recruiter and ATS resume "
                    "analyst. Compare the resume with the job description. "
                    "Return an objective structured analysis. Do not invent "
                    "skills, experience, education, or achievements."
                ),
            },
            {
                "role": "user",
                "content": (
                    "RESUME:\n"
                    f"{resume_text}\n\n"
                    "JOB DESCRIPTION:\n"
                    f"{job_description}"
                ),
            },
        ],
        text_format=ResumeAnalysis,
    )

    if response.output_parsed is None:
        raise RuntimeError(
            "OpenAI returned no structured resume analysis."
        )

    return response.output_parsed


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract readable text from a PDF resume."""
    reader = PdfReader(str(file_path))

    pages: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            pages.append(page_text.strip())

    return "\n\n".join(pages).strip()


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from paragraphs and tables in a DOCX resume."""
    document = Document(str(file_path))

    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if values:
                sections.append(" | ".join(values))

    return "\n".join(sections).strip()


def extract_text_from_txt(file_path: Path) -> str:
    """Read text from a plain-text resume."""
    return file_path.read_text(
        encoding="utf-8",
        errors="ignore",
    ).strip()


def extract_resume_text(file_path: str | Path) -> str:
    """
    Extract resume text from PDF, DOCX, or TXT.

    Raises:
        FileNotFoundError: When the resume file does not exist.
        ValueError: When the file type is unsupported or contains no text.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"Resume file was not found: {path}"
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_RESUME_EXTENSIONS:
        supported = ", ".join(
            sorted(SUPPORTED_RESUME_EXTENSIONS)
        )

        raise ValueError(
            f"Unsupported resume format: {extension}. "
            f"Supported formats: {supported}"
        )

    if extension == ".pdf":
        text = extract_text_from_pdf(path)
    elif extension == ".docx":
        text = extract_text_from_docx(path)
    else:
        text = extract_text_from_txt(path)

    if not text:
        raise ValueError(
            "No readable text was found in the resume."
        )

    return text