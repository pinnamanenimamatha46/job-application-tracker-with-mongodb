from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract readable text from a PDF resume."""
    reader = PdfReader(str(file_path))
    pages: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            pages.append(page_text.strip())

    return "\n\n".join(pages).strip()


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from paragraphs and tables in a DOCX resume."""
    document = Document(str(file_path))
    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            sections.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if values:
                sections.append(" | ".join(values))

    return "\n".join(sections).strip()


def extract_text_from_txt(file_path: Path) -> str:
    """Read text from a plain-text resume."""
    return file_path.read_text(
        encoding="utf-8",
        errors="ignore",
    ).strip()


def extract_resume_text(file_path: str | Path) -> str:
    """
    Extract resume text from PDF, DOCX, or TXT.

    Raises:
        FileNotFoundError: When the resume does not exist.
        ValueError: When the format is unsupported or contains no text.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"Resume file was not found: {path}"
        )

    if not path.is_file():
        raise ValueError(
            f"Resume path is not a file: {path}"
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_RESUME_EXTENSIONS:
        supported = ", ".join(
            sorted(SUPPORTED_RESUME_EXTENSIONS)
        )

        raise ValueError(
            f"Unsupported resume format: {extension}. "
            f"Supported formats: {supported}"
        )

    if extension == ".pdf":
        text = extract_text_from_pdf(path)
    elif extension == ".docx":
        text = extract_text_from_docx(path)
    else:
        text = extract_text_from_txt(path)

    if not text:
        raise ValueError(
            "No readable text was found in the resume."
        )

    return text